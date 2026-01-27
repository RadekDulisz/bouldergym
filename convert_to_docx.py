from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Read the text file
with open('ANALIZA_SYSTEMU.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Split into lines
lines = content.split('\n')

for line in lines:
    # Check if line is a major header (with ===)
    if '=' * 50 in line:
        continue  # Skip separator lines
    
    # Check if line is a section header (numbered like "1. ")
    if re.match(r'^[\d]+\.\s+[A-ZĄĆĘŁŃÓŚŹŻ]', line):
        p = doc.add_heading(line, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Check if line is a subsection (like "1.1. ")
    elif re.match(r'^[\d]+\.[\d]+\.\s+[A-ZĄĆĘŁŃÓŚŹŻ]', line):
        p = doc.add_heading(line, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Check if line is ALL CAPS (like "MODUŁ 1:")
    elif line.strip() and line.strip().isupper() and len(line.strip()) > 5:
        p = doc.add_heading(line.strip(), level=3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Check if line starts with box drawing characters (tables/diagrams)
    elif line.strip().startswith('┌') or line.strip().startswith('│') or line.strip().startswith('└') or line.strip().startswith('├'):
        # Add as code/preformatted text
        p = doc.add_paragraph(line)
        p.style = 'Normal'
        run = p.runs[0]
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    # Check if line starts with bullet points or list markers
    elif line.strip().startswith('•') or line.strip().startswith('-') or line.strip().startswith('✓') or line.strip().startswith('✗'):
        p = doc.add_paragraph(line.strip(), style='List Bullet')
    
    # Empty line
    elif not line.strip():
        doc.add_paragraph()
    
    # Regular text
    else:
        p = doc.add_paragraph(line)
        # Check for special formatting indicators
        if '  AKTOR:' in line or '  NAZWA:' in line or '  WARUNKI' in line:
            run = p.runs[0]
            run.bold = True

# Add metadata
core_properties = doc.core_properties
core_properties.title = 'Analiza Systemu - Boulder Gym Management System'
core_properties.author = 'GitHub Copilot'
core_properties.comments = 'Dokumentacja analizy systemu zarządzania halą wspinaczkową'

# Save the document
doc.save('ANALIZA_SYSTEMU.docx')
print("✅ Dokument ANALIZA_SYSTEMU.docx został utworzony pomyślnie!")
